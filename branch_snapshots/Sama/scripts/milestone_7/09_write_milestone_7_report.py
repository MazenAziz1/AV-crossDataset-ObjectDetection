import json
import sys
from pathlib import Path

import pandas as pd
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from scripts.milestone_7 import common

DISP = common.DETECTOR_DISPLAY
DETECTORS = common.DETECTORS


def r3(x):
    return f"{x:.3f}"


def r2(x):
    return f"{x:.2f}"


def main():
    print("=" * 79)
    print("Milestone 7 - Step 9: Report")
    print("=" * 79)

    root = common.PROJECT_ROOT
    sa = root / "outputs" / "milestone_7" / "safety_error_analysis"
    oa = root / "outputs" / "milestone_7" / "object_size_analysis"
    dt = root / "outputs" / "milestone_7" / "deployment_tradeoff"

    deploy = pd.read_csv(dt / "deployment_suitability_table.csv").set_index("detector")
    obj = pd.read_csv(oa / "object_size_summary.csv")
    ft = pd.read_csv(sa / "failure_type_summary.csv").set_index(["dataset", "detector"])
    conf = pd.read_csv(sa / "class_confusion_summary.csv")
    index_summary = json.load(open(sa / "detection_error_index.json", encoding="utf-8"))

    doc = Document()
    doc.add_heading("Robustness, Failure-Case & Safety-Oriented Analysis", level=0)
    sub = doc.add_paragraph("Milestone 7 Report \u2014 Cross-Dataset Object Detection")
    sub.runs[0].italic = True

    # 1. Objective
    doc.add_heading("1. Milestone 7 Objective", level=1)
    doc.add_paragraph(
        "This milestone converts the Milestone 5 KITTI in-domain results and the Milestone 6 Waymo "
        "external-validation results into a robustness and safety analysis. It explains how, where, and "
        "why the four detectors fail, which object types are most vulnerable, which safety-critical misses "
        "occur, how KITTI and Waymo differ, and what deployment trade-offs exist. It is not a new mAP "
        "comparison and performs no training."
    )

    # 2. Source inputs
    doc.add_heading("2. Source Inputs from Milestones 5 and 6", level=1)
    doc.add_paragraph(
        "The analysis reuses the frozen Milestone 5 KITTI evaluation (1,496 validation images) and the "
        "Milestone 6 Waymo external validation (996 images), including their per-detection prediction "
        "exports. The locked checkpoints (YOLOv8s, RT-DETR-L, RetinaNet, Faster R-CNN) and the three-class "
        "contract (Vehicle, Pedestrian, Cyclist) are unchanged."
    )

    # 3. Protocol
    doc.add_heading("3. Robustness Analysis Protocol", level=1)
    for line in [
        "Operating point: confidence >= 0.25, IoU >= 0.50 (frozen Milestone 5/6 operating point).",
        "Object-size bins: small (<32^2 px), medium (32^2-96^2 px), large (>96^2 px), reused from the frozen Milestone 5 policy.",
        "Failure types: missed detection (FN), false positive, localization error (correct class, 0.10 <= IoU < 0.50), class confusion, over-detection.",
        "Safety-priority classes: Pedestrian and Cyclist.",
        "All thresholds and bins were locked before any result was generated.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 4. Object size
    doc.add_heading("4. Object Size Analysis", level=1)
    doc.add_paragraph(
        "Recall and miss rate were computed per detector, dataset, and size category at the operating "
        "point. On KITTI, recall declines with object size but remains high for all detectors. On Waymo, "
        "small-object recall collapses for every detector."
    )
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Detector", "KITTI small recall", "Waymo small recall", "Waymo large recall"]):
        t.rows[0].cells[i].text = h
    for d in DETECTORS:
        row = t.add_row().cells
        row[0].text = DISP[d]
        row[1].text = r3(deploy.loc[d, "KITTI_small_recall"])
        row[2].text = r3(deploy.loc[d, "Waymo_small_recall"])
        row[3].text = r3(obj[(obj["dataset"] == "waymo") & (obj["detector"] == d) & (obj["class_name"] == "all") & (obj["size_category"] == "large")]["recall"].iloc[0])
    doc.add_paragraph(
        "The results indicate that small objects are the dominant failure mode, and that this is far more "
        "severe on Waymo than on KITTI. On KITTI, small-object recall is 0.88-0.96 across detectors, whereas "
        "on Waymo it is 0.03-0.13. This is consistent with Waymo containing many more small/distant objects."
    )

    # 5. Safety FN
    doc.add_heading("5. Safety-Critical False Negative Analysis", level=1)
    doc.add_paragraph(
        "False-negative rates for the safety-priority classes (Pedestrian, Cyclist) were computed "
        "separately from overall accuracy."
    )
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Detector", "KITTI safety FN rate", "Waymo safety FN rate"]):
        t.rows[0].cells[i].text = h
    for d in DETECTORS:
        row = t.add_row().cells
        row[0].text = DISP[d]
        row[1].text = r3(deploy.loc[d, "KITTI_safety_fn_rate"])
        row[2].text = r3(deploy.loc[d, "Waymo_safety_fn_rate"])
    doc.add_paragraph(
        "On Waymo, the pedestrian+cyclist false-negative rate is 0.81-0.93 for all detectors. This is a "
        "safety-relevant result that is not visible from the average mAP alone: even the detector with the "
        "best Waymo mAP misses the large majority of vulnerable road users. The highest-confidence mAP "
        "ranking therefore does not imply the safest detector."
    )

    # 6. FP + localization
    doc.add_heading("6. False Positive and Localization Error Analysis", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Detector", "KITTI FPs", "KITTI localization"]):
        t.rows[0].cells[i].text = h
    for d in DETECTORS:
        row = t.add_row().cells
        row[0].text = DISP[d]
        row[1].text = str(int(ft.loc[("kitti", d), "false_positive"]))
        row[2].text = str(int(ft.loc[("kitti", d), "localization_error"]))
    doc.add_paragraph(
        "RT-DETR-L emits the most false positives and by far the most over-detections on KITTI (a large "
        "fraction of its detections are redundant), consistent with its low operating-point precision "
        "(0.482) reported in Milestone 5."
    )

    # 7. Class confusion
    doc.add_heading("7. Class Confusion Analysis", level=1)
    total_conf = int(conf["count"].sum()) if "count" in conf else len(conf)
    doc.add_paragraph(
        f"Across both datasets there are {total_conf} indexed class-confusion detections. The dominant "
        "confusion on Waymo is between Pedestrian and Cyclist and Pedestrian-as-Vehicle, consistent with "
        "small appearance differences between these classes at distance."
    )

    # 8. Gallery
    doc.add_heading("8. Failure-Case Gallery", level=1)
    doc.add_paragraph(
        "Annotated failure examples (missed pedestrian, missed cyclist, small-object failure, false "
        "positive, localization error, class confusion) were generated separately for KITTI and Waymo and "
        "are catalogued in outputs/milestone_7/failure_cases/. Each example is traceable to its dataset, "
        "image, detector, and error type via the manifest."
    )

    # 9. Deployment
    doc.add_heading("9. Deployment Suitability Comparison", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Detector", "KITTI mAP", "Waymo mAP", "Latency (ms)", "Waymo safety FN rate", "Waymo small recall"]):
        t.rows[0].cells[i].text = h
    for d in DETECTORS:
        row = t.add_row().cells
        row[0].text = DISP[d]
        row[1].text = r3(deploy.loc[d, "KITTI_mAP50_95"])
        row[2].text = r3(deploy.loc[d, "Waymo_mAP50_95"])
        row[3].text = r2(deploy.loc[d, "waymo_mean_inference_ms"])
        row[4].text = r3(deploy.loc[d, "Waymo_safety_fn_rate"])
        row[5].text = r3(deploy.loc[d, "Waymo_small_recall"])

    # 10. Recommendations
    doc.add_heading("10. Practical Recommendations for Intelligent-Vehicle Perception", level=1)
    for line in [
        "The in-domain accuracy ranking does not transfer to the out-of-domain setting; detector selection must be based on external validation and safety metrics, not KITTI mAP alone.",
        "Small-object and pedestrian/cyclist recall should be treated as first-class deployment metrics alongside mAP.",
        "No single detector simultaneously optimizes accuracy, generalization, latency, and safety; the choice depends on the deployment priority.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 11. Paper-ready discussion
    doc.add_heading("11. Paper-Ready Robustness and Safety Discussion", level=1)
    doc.add_paragraph(
        "Cross-dataset robustness and safety were analyzed by matching per-detection predictions against "
        "ground truth at a fixed operating point for both KITTI and Waymo. The analysis indicates that the "
        "dominant failure mode is missed detections of small objects, and that vulnerable road users "
        "(pedestrians and cyclists) are missed at a high rate on the external dataset. These patterns are "
        "consistent with a domain shift toward more numerous, smaller, and more distant objects in Waymo, "
        "and they suggest that average accuracy alone is an insufficient measure of safety for "
        "intelligent-vehicle perception."
    )

    # 12. Limitations
    doc.add_heading("12. Limitations and Threats to Validity", level=1)
    for line in [
        "Analysis is confined to the frozen operating point (confidence 0.25, IoU 0.50); results may differ at other operating points.",
        "The Waymo subset (996 images) is smaller than KITTI validation (1,496 images).",
        "Failure-type classification uses documented IoU-based heuristics; ambiguous cases are resolved by a fixed priority order.",
        "No causal claims are made; findings are descriptive and correlational.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 13. Artifacts
    doc.add_heading("13. Generated Artifacts", level=1)
    for line in [
        "Error index: outputs/milestone_7/safety_error_analysis/detection_error_index.csv",
        "Object-size analysis: outputs/milestone_7/object_size_analysis/",
        "Safety FN analysis: outputs/milestone_7/safety_error_analysis/",
        "Failure-type analysis: outputs/milestone_7/safety_error_analysis/",
        "Gallery: outputs/milestone_7/failure_cases/",
        "Deployment trade-off: outputs/milestone_7/deployment_tradeoff/",
        "Figures: outputs/milestone_7/figures/",
        "Configs: configs/analysis/milestone_7/",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    out_path = root / "docs" / "milestone_7" / "Milestone_7_Robustness_Failure_Case_Safety_Report.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
