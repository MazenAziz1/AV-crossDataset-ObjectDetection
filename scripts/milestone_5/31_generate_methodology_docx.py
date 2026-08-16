import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DETECTORS = ["yolo", "faster_rcnn", "retinanet", "rtdetr"]
ORDER = ["yolo", "rtdetr", "retinanet", "faster_rcnn"]
NAMES = {
    "yolo": "YOLOv8s",
    "rtdetr": "RT-DETR-L",
    "retinanet": "RetinaNet",
    "faster_rcnn": "Faster R-CNN",
}

# Training facts (recorded from Kaggle session logs; not stored in JSONs)
TRAINING = {
    "yolo": {"epochs": "200 / 200", "best": "final", "stop": "completed", "time": "6.33 h"},
    "rtdetr": {"epochs": "131 / 200", "best": "epoch 111", "stop": "early stop (patience 20)", "time": "multi-session"},
    "retinanet": {"epochs": "101 / 200", "best": "epoch 81", "stop": "early stop (patience 20)", "time": "multi-session"},
    "faster_rcnn": {"epochs": "96 / 200", "best": "pre-resume (mAP 0.5905)", "stop": "early stop (patience 20)", "time": "multi-session"},
}


def r4(x):
    return f"{x:.4f}"


def r2(x):
    return f"{x:.2f}"


def main():
    root = Path(__file__).resolve().parents[2]
    metrics_dir = root / "outputs" / "milestone_5" / "metrics" / "kitti_validation"
    benchmarks_dir = root / "outputs" / "milestone_5" / "benchmarks"
    out_path = root / "docs" / "milestone_5" / "model_training_and_kitti_evaluation_draft.docx"

    metrics = {d: json.load(open(metrics_dir / f"{d}_metrics.json")) for d in DETECTORS}
    benchmarks = {d: json.load(open(benchmarks_dir / f"{d}_benchmark.json")) for d in DETECTORS}

    doc = Document()

    # Title
    doc.add_heading("Model Training and KITTI Evaluation", level=0)
    sub = doc.add_paragraph("Draft Report \u2014 Milestones 4 + 5 (Cross-Dataset Object Detection)")
    sub.runs[0].italic = True

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Four object detectors (YOLOv8s, RT-DETR-L, RetinaNet, and Faster R-CNN) were trained on the "
        "KITTI training partition (5,985 images) and evaluated on the KITTI validation partition "
        "(1,496 images) under a frozen experimental protocol. Training ran on Kaggle GPU notebooks "
        "(Tesla T4) across two compute slots; evaluation, benchmarking, and checkpoint locking ran locally."
    )
    for line in [
        "Input size: 640 x 640 (letterboxed)",
        "Target classes: Vehicle, Pedestrian, Cyclist",
        "Random seed: 42 (deterministic mode)",
        "Primary metric: COCO mAP@[0.50:0.95]",
        "Early stopping: patience = 20 epochs on mAP@[0.50:0.95]",
        "Target epochs: 200",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 2. Methodology
    doc.add_heading("2. Methodology", level=1)
    doc.add_heading("2.1 Evaluation protocol", level=2)
    doc.add_paragraph(
        "The frozen evaluation protocol computes COCO Average Precision (mAP@[0.50:0.95], AP50, AP75), "
        "per-class AP, and scale-specific AP (small/medium/large). KITTI DontCare regions are suppressed "
        "at IoU >= 0.50 before evaluation."
    )
    doc.add_paragraph(
        "For the torchvision-based detectors, the training pipeline applied ImageNet normalization through "
        "Albumentations in addition to the normalization performed internally by the torchvision detection "
        "models. Consequently, evaluation was performed using the same preprocessing configuration employed "
        "during training to avoid a train-evaluation distribution mismatch. The Ultralytics-based YOLOv8s "
        "and RT-DETR models did not exhibit this issue."
    )
    doc.add_heading("2.2 Kaggle compute allocation", level=2)
    doc.add_paragraph(
        "Slot A ran YOLOv8s, Faster R-CNN, and RetinaNet. Slot B ran RT-DETR-L as an interruption-expected "
        "run across a multi-session resume chain. All models used an effective batch size of 32 (RT-DETR-L: 16), a 10.5-hour "
        "runtime guard, and checkpoint/resume-state packaging between sessions."
    )

    # 3. Training summary
    doc.add_heading("3. Training Summary", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Detector", "Epochs", "Best checkpoint", "Exit reason", "Training time"]):
        hdr[i].text = h
    for d in ORDER:
        f = TRAINING[d]
        row = t.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = f["epochs"]
        row[2].text = f["best"]
        row[3].text = f["stop"]
        row[4].text = f["time"]

    # 4. Results
    doc.add_heading("4. Results (KITTI validation, local protocol)", level=1)

    doc.add_heading("4.1 Accuracy", level=2)
    acc = doc.add_table(rows=1, cols=7)
    acc.style = "Light Grid Accent 1"
    ah = acc.rows[0].cells
    for i, h in enumerate(["Detector", "mAP@0.50:0.95", "mAP@0.50", "mAP@0.75", "Vehicle AP", "Pedestrian AP", "Cyclist AP"]):
        ah[i].text = h
    for d in ORDER:
        m = metrics[d]
        pc = m["per_class"]
        row = acc.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = r4(m["metrics"]["mAP_50_95"])
        row[2].text = r4(m["metrics"]["mAP_50"])
        row[3].text = r4(m["metrics"]["mAP_75"])
        row[4].text = r4(pc["1"]["AP_50_95"])
        row[5].text = r4(pc["2"]["AP_50_95"])
        row[6].text = r4(pc["3"]["AP_50_95"])

    doc.add_heading("4.2 Efficiency", level=2)
    eff = doc.add_table(rows=1, cols=5)
    eff.style = "Light Grid Accent 1"
    eh = eff.rows[0].cells
    for i, h in enumerate(["Detector", "Parameters", "Checkpoint size", "Latency (ms)", "FPS (local)"]):
        eh[i].text = h
    for d in ORDER:
        b = benchmarks[d]
        row = eff.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = f"{b['parameter_count']:,}"
        row[2].text = f"{b['checkpoint_size_mb']} MB"
        row[3].text = r2(b["latency_total_ms"])
        row[4].text = r2(b["frames_per_second"])

    doc.add_heading("4.3 Operating Point (confidence >= 0.25, IoU >= 0.50)", level=2)
    op = doc.add_table(rows=1, cols=5)
    op.style = "Light Grid Accent 1"
    oh = op.rows[0].cells
    for i, h in enumerate(["Detector", "Precision", "Recall", "F1", "False positives/image"]):
        oh[i].text = h
    for d in ORDER:
        o = metrics[d].get("operating_point", {})
        row = op.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = r4(o.get("precision", 0.0))
        row[2].text = r4(o.get("recall", 0.0))
        row[3].text = r4(o.get("f1_score", 0.0))
        row[4].text = r4(o.get("false_positives_per_image", 0.0))

    # 5. Findings
    doc.add_heading("5. Key Findings", level=1)
    for line in [
        "YOLOv8s dominates every axis: highest accuracy (mAP@0.50:0.95 = 0.690) at the lowest parameter "
        "count (11.1M) and ~5-8x faster than the other detectors (94.77 FPS).",
        "Accuracy ranking: YOLOv8s > RT-DETR-L > Faster R-CNN > RetinaNet.",
        "Pedestrian is the weakest class across all four detectors (0.392-0.530 AP), consistent with the "
        "class's small object size and limited training instances.",
        "The two heaviest detectors carry a large efficiency penalty: Faster R-CNN (41.3M params, 11.24 FPS, "
        "315.68 MB) and RetinaNet (36.4M params, 13.70 FPS, 278.0 MB).",
        "RT-DETR-L offers the second-best accuracy but remains far slower than YOLOv8s at deployment scale.",
        "At the fixed 0.25 operating point, RT-DETR-L drops to 0.482 precision (5.43 false positives/image) "
        "because its DETR decoder emits many low-confidence detections.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 6. Boundary
    doc.add_heading("6. Milestone 6 Boundary", level=1)
    doc.add_paragraph(
        "Waymo cross-domain evaluation, cross-domain degradation measurement, and domain-shift analysis are "
        "deferred to Milestone 6 and are explicitly out of scope for this report."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
