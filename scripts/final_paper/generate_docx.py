"""Generate the DOCX version of the finalized LaTeX paper.

This script is a faithful conversion of ``docs/final_paper/main.tex`` (and its
``sections/``, ``tables/``, and ``figures/``) into a professional Word document.
The LaTeX paper remains the authoritative source; this script only *reads* it.

Content fidelity:
  * Sections / subsections / subsubsections are preserved with explicit numbering
    (Abstract, 1. Introduction, ..., 9. Conclusion, References).
  * Body text, inline math (rendered as Unicode), bold/italic/monospace runs,
    and citation/table/figure cross-references are converted from the LaTeX source.
  * Figures are inserted in source order with their LaTeX captions.  PDF data
    figures are rasterized by re-running the same matplotlib code (and source CSVs)
    that produced the PDFs, because python-docx cannot embed PDF; the two existing
    PNG pipeline diagrams are embedded directly.
  * Tables are converted into real Word tables (booktabs look), preserving every
    value from the generated ``tables/*.tex`` files.
  * References are produced from ``references.bib`` in first-appearance order
    (matching the ``unsrt`` BibTeX style used by the paper).

Usage::

    python scripts/final_paper/generate_docx.py
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "docs" / "final_paper"
SECTIONS_DIR = PAPER / "sections"
TABLES_DIR = PAPER / "tables"
FIGURES_DIR = PAPER / "figures"
OUTPUT = PAPER / "AV_Cross_Dataset_Object_Detection_Research_Paper.docx"

SECTION_ORDER = [
    "introduction",
    "related_work",
    "datasets",
    "methodology",
    "experimental_setup",
    "results",
    "discussion",
    "threats_to_validity",
    "conclusion",
]

# --------------------------------------------------------------------------- #
# Unicode math / accent helpers
# --------------------------------------------------------------------------- #
_SUB = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")
_SUP = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")

ACCENT_MAP = {
    "'": {"a": "á", "e": "é", "i": "í", "o": "ó", "u": "ú", "y": "ý", "c": "ć"},
    "`": {"a": "à", "e": "è", "i": "ì", "o": "ò", "u": "ù"},
    '"': {"a": "ä", "e": "ë", "i": "ï", "o": "ö", "u": "ü", "y": "ÿ"},
    "^": {"a": "â", "e": "ê", "i": "î", "o": "ô", "u": "û"},
    "~": {"a": "ã", "n": "ñ", "o": "õ"},
}


def convert_math(m: str) -> str:
    m = m.replace("{,}", "").replace("{:}", ":").replace("{}", "")
    m = m.replace("{", "").replace("}", "")
    m = m.replace(r"\times", "\u00d7")
    m = m.replace(r"\geq", "\u2265").replace(r"\leq", "\u2264")
    m = m.replace(r"\min", "min").replace(r"\max", "max")
    m = m.replace(r"\%", "%").replace(r"\ ", " ")
    m = m.replace(r"\,", "")
    m = re.sub(r"_\{([^}]*)\}", lambda mm: mm.group(1).translate(_SUB), m)
    m = re.sub(r"\^\{([^}]*)\}", lambda mm: mm.group(1).translate(_SUP), m)
    m = re.sub(r"_([0-9A-Za-z])", lambda mm: mm.group(1).translate(_SUB), m)
    m = re.sub(r"\^([0-9A-Za-z])", lambda mm: mm.group(1).translate(_SUP), m)
    return m


def decode_accents(s: str) -> str:
    def repl(mm: re.Match) -> str:
        acc = mm.group(1)[1]
        return ACCENT_MAP.get(acc, {}).get(mm.group(2), mm.group(2))

    return re.sub(r"\{(\\.)([a-zA-Z])\}", repl, s)


# --------------------------------------------------------------------------- #
# Braced-command extraction
# --------------------------------------------------------------------------- #
def _find_matching(s: str, start: int) -> int:
    depth = 0
    i = start
    while i < len(s):
        if s[i] == "\\":
            i += 2
            continue
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return len(s)


def extract_command(s: str, cmd: str) -> str:
    i = s.index(cmd)
    b = s.index("{", i + len(cmd))
    e = _find_matching(s, b)
    return s[b + 1 : e]


# --------------------------------------------------------------------------- #
# Inline parser: LaTeX -> list of (text, bold, italic, code) runs
# --------------------------------------------------------------------------- #
_CMD_STYLES = [
    ("\\textbf", (True, False, False)),
    ("\\emph", (False, True, False)),
    ("\\textit", (False, True, False)),
    ("\\texttt", (False, False, True)),
]
_ESCAPES = {"\\%": "%", "\\_": "_", "\\&": "&", "\\#": "#", "\\{": "{", "\\}": "}", "\\$": "$"}


def parse_inline(s: str, ref_map, cite_map):
    s = s.replace("{,}", ",").replace("{:}", ":")
    runs = []
    i = 0
    n = len(s)
    buf = ""

    def flush(b=False, it=False, code=False):
        nonlocal buf
        if buf:
            runs.append((buf, b, it, code))
            buf = ""

    while i < n:
        c = s[i]
        if c == "$":
            j = s.index("$", i + 1)
            buf += convert_math(s[i + 1 : j])
            i = j + 1
            continue
        if c == "\\":
            handled = False
            for cmd, (b, it, code) in _CMD_STYLES:
                if s.startswith(cmd + "{", i):
                    flush()
                    close = _find_matching(s, i + len(cmd))
                    inner = s[i + len(cmd) + 1 : close]
                    for (t, bb, ii, cc) in parse_inline(inner, ref_map, cite_map):
                        runs.append((t, b or bb, it or ii, code or cc))
                    i = close + 1
                    handled = True
                    break
            if handled:
                continue
            if s.startswith("\\cite{", i):
                flush()
                close = _find_matching(s, i + 5)
                keys = [k.strip() for k in s[i + 6 : close].split(",")]
                nums = [str(cite_map[k]) for k in keys if k]
                runs.append(("[" + ", ".join(nums) + "]", False, False, False))
                i = close + 1
                continue
            if s.startswith("\\ref{", i):
                flush()
                close = _find_matching(s, i + 4)
                key = s[i + 5 : close].strip()
                runs.append((ref_map.get(key, "?"), False, False, False))
                i = close + 1
                continue
            if s.startswith("\\label{", i):
                flush()
                close = _find_matching(s, i + 6)
                i = close + 1
                continue
            matched = False
            for e, r in _ESCAPES.items():
                if s.startswith(e, i):
                    buf += r
                    i += len(e)
                    matched = True
                    break
            if matched:
                continue
            if s.startswith("\\ ", i):
                buf += " "
                i += 2
                continue
            if s.startswith("\\textasciitilde{}", i):
                buf += "~"
                i += len("\\textasciitilde{}")
                continue
            buf += "\\"
            i += 1
            continue
        if s.startswith("---", i):
            buf += "\u2014"
            i += 3
            continue
        if s.startswith("--", i):
            buf += "\u2013"
            i += 2
            continue
        if c == "~":
            buf += "\u00a0"
            i += 1
            continue
        buf += c
        i += 1
    flush()
    return runs


# --------------------------------------------------------------------------- #
# Block parser: a section .tex file -> list of blocks
# --------------------------------------------------------------------------- #
_BLOCK_STARTS = (
    "\\section", "\\subsection", "\\subsubsection", "\\begin{", "\\end{",
    "\\label", "\\input", "\\keywords", "\\item",
)


def _is_block_start(line: str) -> bool:
    ls = line.lstrip()
    if not ls:
        return True
    return ls.startswith(_BLOCK_STARTS)


def parse_section(path: Path):
    raw = path.read_text(encoding="utf-8").split("\n")
    lines = [ln for ln in raw if not ln.lstrip().startswith("%")]
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("\\begin{figure}"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("\\end{figure}"):
                buf.append(lines[i])
                i += 1
            i += 1
            content = "\n".join(buf)
            m_img = re.search(r"\\includegraphics[^\{]*\{([^}]*)\}", content)
            img = m_img.group(1) if m_img else None
            cap = None
            if "\\caption" in content:
                cap = extract_command(content, "\\caption")
            lab = None
            if "\\label" in content:
                lab = extract_command(content, "\\label")
            blocks.append(("figure", img, cap, lab))
            continue
        if line.startswith("\\begin{enumerate}"):
            items = []
            cur = None
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("\\end{enumerate}"):
                ln = lines[i].strip()
                if ln.startswith("\\item"):
                    if cur is not None:
                        items.append(cur)
                    cur = ln[len("\\item") :].strip()
                else:
                    if cur is not None:
                        cur += " " + ln
                i += 1
            if cur is not None:
                items.append(cur)
            i += 1
            blocks.append(("enumerate", items))
            continue
        if line.startswith("\\begin{abstract}"):
            txt = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("\\end{abstract}"):
                txt.append(lines[i].strip())
                i += 1
            i += 1
            blocks.append(("abstract", " ".join(txt)))
            continue
        if line.startswith("\\keywords{"):
            blocks.append(("keywords", extract_command(line, "\\keywords")))
            i += 1
            continue
        for kind, cmd in [("section", "\\section"), ("subsection", "\\subsection"), ("subsubsection", "\\subsubsection")]:
            if line.startswith(cmd + "{"):
                blocks.append((kind, extract_command(line, cmd)))
                i += 1
                break
        else:
            if line.startswith("\\label{"):
                blocks.append(("label", extract_command(line, "\\label")))
                i += 1
                continue
            if line.startswith("\\input{"):
                blocks.append(("table", extract_command(line, "\\input")))
                i += 1
                continue
            para = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i]):
                para.append(lines[i].strip())
                i += 1
            blocks.append(("para", " ".join(para)))
            continue
    return blocks


# --------------------------------------------------------------------------- #
# Table parser: a tables/*.tex file -> caption, label, rows (with colspans)
# --------------------------------------------------------------------------- #
def parse_table_tex(path: Path):
    txt = path.read_text(encoding="utf-8")
    cap = extract_command(txt, "\\caption") if "\\caption" in txt else ""
    lab = extract_command(txt, "\\label") if "\\label" in txt else ""
    m_spec = re.search(r"\\begin\{tabular\}\{([^}]*)\}", txt)
    ncols = len(re.sub(r"[^lcrp]", "", m_spec.group(1))) if m_spec else 0
    m_body = re.search(r"\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}", txt, re.S)
    body = m_body.group(1) if m_body else ""
    header_rows = []
    data_rows = []
    in_header = True
    for raw_line in body.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("\\toprule") or line.startswith("\\midrule") or line.startswith("\\bottomrule"):
            if line.startswith("\\midrule"):
                in_header = False
            continue
        if line == "\\addlinespace":
            continue
        line = re.sub(r"\\\\$", "", line)
        cells = _split_cells(line)
        if in_header:
            header_rows.append(cells)
        else:
            data_rows.append(cells)
    return cap, lab, ncols, header_rows, data_rows


def _split_cells(line: str):
    cells = []
    for seg in line.split("&"):
        seg = seg.strip()
        if seg.startswith("\\multicolumn"):
            m = re.match(r"\\multicolumn\{(\d+)\}\{[^}]*\}\{", seg)
            close = _find_matching(seg, len(m.group(0)) - 1)
            text = seg[len(m.group(0)) : close].strip()
            cells.append(("multi", int(m.group(1)), text))
        else:
            cells.append(seg)
    return cells


# --------------------------------------------------------------------------- #
# BibTeX parsing / reference formatting
# --------------------------------------------------------------------------- #
def parse_bib(path: Path):
    txt = path.read_text(encoding="utf-8")
    entries = {}
    for m in re.finditer(r"@(\w+)\{([^,]+),\s*(.*?)\n\}", txt, re.S):
        key = m.group(2)
        body = m.group(3)
        fields = {}
        cur_key = None
        cur_val = []
        for line in body.split("\n"):
            line = line.rstrip().rstrip(",").strip()
            if not line:
                continue
            fm = re.match(r"(\w+)\s*=\s*(.*)$", line)
            if fm and "{" in fm.group(2):
                if cur_key:
                    fields[cur_key] = " ".join(cur_val)
                cur_key = fm.group(1)
                cur_val = [fm.group(2).strip()]
            else:
                cur_val.append(line)
        if cur_key:
            fields[cur_key] = " ".join(cur_val)
        for k, v in fields.items():
            v = v.strip()
            if v.startswith("{") and v.endswith("}"):
                v = v[1:-1]
            elif v.startswith('"') and v.endswith('"'):
                v = v[1:-1]
            fields[k] = v
        entries[key] = fields
    return entries


def _initials(first: str) -> str:
    toks = first.replace("~", " ").split()
    out = []
    for t in toks:
        if "-" in t:
            out.append("-".join(w[0].upper() + "." for w in t.split("-") if w))
        else:
            out.append(t[0].upper() + ".")
    return " ".join(out)


def format_authors(author: str) -> str:
    author = decode_accents(author)
    parts = [p.strip() for p in re.split(r"\s+and\s+", author)]
    named = []
    etal = False
    for p in parts:
        if p.lower() == "others":
            etal = True
            continue
        if "," in p:
            last, first = p.split(",", 1)
            named.append(_initials(first.strip()) + " " + last.strip())
        else:
            named.append(p)
    if etal:
        named.append("et al.")
    if len(named) == 1:
        return named[0]
    if named and named[-1] == "et al.":
        return ", ".join(named[:-1]) + ", et al."
    return ", ".join(named[:-1]) + ", and " + named[-1]


def strip_braces(s: str) -> str:
    return decode_accents(s.replace("{", "").replace("}", ""))


def format_reference(entry: dict) -> str:
    authors = format_authors(entry.get("author", ""))
    title = strip_braces(entry.get("title", ""))
    booktitle = strip_braces(entry.get("booktitle", ""))
    year = entry.get("year", "")
    return f'{authors}, "{title}," in {booktitle}, {year}.'


# --------------------------------------------------------------------------- #
# Figure rasterization (PDF data figures -> PNG), mirroring generate_figures.py
# --------------------------------------------------------------------------- #
M6_TABLES = ROOT / "outputs" / "milestone_6" / "generalization_analysis" / "tables"
M7 = ROOT / "outputs" / "milestone_7"

DETECTOR_ORDER = ["yolo", "rtdetr", "faster_rcnn", "retinanet"]
DISPLAY = {
    "yolo": "YOLOv8s",
    "rtdetr": "RT-DETR-L",
    "faster_rcnn": "Faster R-CNN",
    "retinanet": "RetinaNet",
}
COLORS = {
    "yolo": "#4C72B0",
    "rtdetr": "#DD8452",
    "faster_rcnn": "#C44E52",
    "retinanet": "#55A868",
}
DATASET_LABEL = {"kitti": "KITTI validation", "waymo": "Waymo external"}


def _load(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path)
    df = df[df["detector"].isin(DETECTOR_ORDER)]
    df["detector"] = pd.Categorical(df["detector"], categories=DETECTOR_ORDER, ordered=True)
    return df.sort_values("detector").reset_index(drop=True)


def _xticklabels() -> list:
    return [DISPLAY[d] for d in DETECTOR_ORDER]


def _plot_compare(path: Path) -> plt.Figure:
    df = _load(path)
    x = np.arange(len(df))
    w = 0.38
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ax.bar(x - w / 2, df["KITTI_mAP50_95"], w, label="KITTI validation", color="#8DA0CB")
    ax.bar(x + w / 2, df["Waymo_mAP50_95"], w, label="Waymo external", color="#FC8D62")
    ax.set_ylabel("mAP@0.50:0.95")
    ax.set_xticks(x)
    ax.set_xticklabels(_xticklabels())
    ax.set_ylim(0, df["KITTI_mAP50_95"].max() * 1.12)
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _plot_ratio(path: Path) -> plt.Figure:
    df = _load(path)
    x = np.arange(len(df))
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ax.bar(x, df["mAP50_95_generalization_ratio"], color=[COLORS[d] for d in df["detector"]])
    ax.axhline(1.0, color="gray", linestyle="--", linewidth=1, label="No drop (ratio = 1.0)")
    ax.set_ylabel("Generalization ratio")
    ax.set_xticks(x)
    ax.set_xticklabels(_xticklabels())
    ax.set_ylim(0, df["mAP50_95_generalization_ratio"].max() * 1.2)
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _plot_class_degrad(path: Path) -> plt.Figure:
    df = _load(path)
    classes = ["Vehicle", "Pedestrian", "Cyclist"]
    x = np.arange(len(classes))
    w = 0.2
    fig, ax = plt.subplots(figsize=(7, 3.8))
    for i, det in enumerate(DETECTOR_ORDER):
        sub = df[df["detector"] == det].set_index("class_name").loc[classes]
        ax.bar(x + (i - 1.5) * w, sub["absolute_drop"], w, label=DISPLAY[det], color=COLORS[det])
    ax.set_ylabel("Absolute drop in AP@0.50:0.95")
    ax.set_xticks(x)
    ax.set_xticklabels(classes)
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _plot_size_recall(path: Path) -> plt.Figure:
    df = _load(path)
    df = df[df["class_name"] == "all"]
    sizes = ["small", "medium", "large"]
    labels = ["Small", "Medium", "Large"]
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.6), sharey=True)
    x = np.arange(3)
    w = 0.2
    for ax, ds in zip(axes, ["kitti", "waymo"]):
        sub = df[df["dataset"] == ds]
        for i, det in enumerate(DETECTOR_ORDER):
            r = sub[sub["detector"] == det].set_index("size_category").loc[sizes]
            ax.bar(x + (i - 1.5) * w, r["recall"], w, label=DISPLAY[det], color=COLORS[det])
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_title(DATASET_LABEL[ds])
        ax.set_ylim(0, 1.0)
        ax.grid(axis="y", alpha=0.3)
    axes[0].set_ylabel("Recall")
    axes[0].legend(fontsize=8)
    fig.tight_layout()
    return fig


def _plot_fn_rate(path: Path) -> plt.Figure:
    df = _load(path)
    df = df[df["class_name"] == "pedestrian+cyclist"]
    agg = df.groupby(["dataset", "detector"]).agg(tp=("tp", "sum"), fn=("fn", "sum")).reset_index()
    agg["fn_rate"] = agg["fn"] / (agg["tp"] + agg["fn"])
    x = np.arange(len(DETECTOR_ORDER))
    w = 0.38
    fig, ax = plt.subplots(figsize=(7, 3.6))
    for offset, ds in zip([-w / 2, w / 2], ["kitti", "waymo"]):
        sub = agg[agg["dataset"] == ds].set_index("detector").loc[DETECTOR_ORDER]
        ax.bar(x + offset, sub["fn_rate"], w, label=DATASET_LABEL[ds])
    ax.set_xticks(x)
    ax.set_xticklabels(_xticklabels())
    ax.set_ylabel("Pedestrian + cyclist FN rate")
    ax.set_ylim(0, 1.0)
    ax.legend()
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _plot_deployment(path: Path) -> plt.Figure:
    df = _load(path)
    fig, ax = plt.subplots(figsize=(6.6, 3.8))
    for _, r in df.iterrows():
        ax.scatter(
            r["waymo_mean_inference_ms"],
            r["Waymo_mAP50_95"],
            s=140,
            color=COLORS[r["detector"]],
            label=DISPLAY[r["detector"]],
        )
    ax.set_xlabel("Mean inference time (ms/image, Waymo)")
    ax.set_ylabel("Waymo mAP@0.50:0.95")
    ax.legend(fontsize=8)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return fig


FIGURE_PLOTTERS = {
    "kitti_vs_waymo_map50_95.pdf": ("kitti_vs_waymo_comparison.csv", _plot_compare),
    "generalization_ratio_map50_95.pdf": ("kitti_vs_waymo_comparison.csv", _plot_ratio),
    "class_wise_degradation.pdf": ("class_wise_degradation.csv", _plot_class_degrad),
    "object_size_recall.pdf": ("object_size_summary.csv", _plot_size_recall),
    "pedestrian_cyclist_fn_rate.pdf": ("safety_false_negative_summary.csv", _plot_fn_rate),
    "deployment_tradeoff.pdf": ("deployment_suitability_table.csv", _plot_deployment),
}


def rasterize_figures(tmpdir: Path) -> dict:
    """Return a mapping from figure file name (in figures/) -> PNG path on disk."""
    png_map = {}
    for pdf_name, (csv_name, plotter) in FIGURE_PLOTTERS.items():
        if csv_name == "object_size_summary.csv":
            csv_path = M7 / "object_size_analysis" / csv_name
        elif csv_name == "safety_false_negative_summary.csv":
            csv_path = M7 / "safety_error_analysis" / csv_name
        elif csv_name == "deployment_suitability_table.csv":
            csv_path = M7 / "deployment_tradeoff" / csv_name
        else:
            csv_path = M6_TABLES / csv_name
        fig = plotter(csv_path)
        png = tmpdir / (pdf_name[:-4] + ".png")
        fig.savefig(png, dpi=200)
        plt.close(fig)
        png_map[pdf_name] = png
    # pipeline PNGs are used directly
    png_map["pipeline.png"] = FIGURES_DIR / "pipeline.png"
    png_map["partition_roles.png"] = FIGURES_DIR / "partition_roles.png"
    return png_map


# --------------------------------------------------------------------------- #
# DOCX styling helpers
# --------------------------------------------------------------------------- #
def _set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _shade_tc(tc, fill: str):
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, val in [("top", "single"), ("bottom", "single"),
                      ("left", "none"), ("right", "none"),
                      ("insideH", "none"), ("insideV", "none")]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), val)
        if val == "single":
            el.set(qn("w:sz"), "8")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _tc_bottom_border(tc):
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    tcBorders.append(bottom)


def _set_cell_text(cell, text, bold, align):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def _set_gridspan(tc, span):
    tcPr = tc.get_or_add_tcPr()
    for gs in tcPr.findall(qn("w:gridSpan")):
        tcPr.remove(gs)
    if span > 1:
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(span))
        tcPr.append(gs)


def _remove_next_tcs(tc, n):
    tr = tc.getparent()
    removed = 0
    el = tc.getnext()
    while el is not None and removed < n:
        nxt = el.getnext()
        if el.tag == qn("w:tc"):
            tr.remove(el)
            removed += 1
        el = nxt


def _set_vmerge(tc, restart):
    tcPr = tc.get_or_add_tcPr()
    vm = OxmlElement("w:vMerge")
    vm.set(qn("w:val"), "restart" if restart else "continue")
    tcPr.append(vm)


def add_table(doc, caption, label, ncols, header_rows, data_rows, ref_map, cite_map):
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(8)
    cap_p.paragraph_format.space_after = Pt(4)
    run = cap_p.add_run("Table " + label + ". ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    for (t, b, it, code) in parse_inline(caption, ref_map, cite_map):
        r = cap_p.add_run(t)
        r.bold = b
        r.italic = it
        r.font.name = "Consolas" if code else "Times New Roman"
        r.font.size = Pt(10)

    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = 1  # center
    _set_table_borders(table)

    header_tc_lists = []

    def fill_row(cells, bold):
        row = table.add_row()
        tcs = list(row._tr.tc_lst)  # snapshot of this row's tc elements
        idx = 0
        for c in cells:
            if isinstance(c, tuple):
                _, span, text = c
                cell = _Cell(tcs[idx], table)
                _set_cell_text(cell, text, bold, WD_ALIGN_PARAGRAPH.CENTER)
                if span > 1:
                    _set_gridspan(tcs[idx], span)
                    _remove_next_tcs(tcs[idx], span - 1)
                idx += span
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_text(_Cell(tcs[idx], table), c, bold, align)
                idx += 1
        return tcs

    for r in header_rows:
        header_tc_lists.append(fill_row(r, True))
    for r in data_rows:
        fill_row(r, False)

    # header styling: shade + midrule (bottom border of last header row)
    for ri, tcs in enumerate(header_tc_lists):
        for tc in tcs:
            _shade_tc(tc, "D9E2F3")
        if ri == len(header_tc_lists) - 1:
            for tc in tcs:
                _tc_bottom_border(tc)

    # vertical merge for a two-level header whose second row starts empty
    if len(header_tc_lists) == 2 and header_rows[1] and header_rows[1][0] == "":
        _set_vmerge(header_tc_lists[0][0], restart=True)
        _set_vmerge(header_tc_lists[1][0], restart=False)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# --------------------------------------------------------------------------- #
# Main assembly
# --------------------------------------------------------------------------- #
def main():
    tmpdir = Path(tempfile.mkdtemp(prefix="av_docx_"))
    png_map = rasterize_figures(tmpdir)

    all_blocks = []
    for name in SECTION_ORDER:
        all_blocks.append((name, parse_section(SECTIONS_DIR / f"{name}.tex")))
    abstract_blocks = parse_section(SECTIONS_DIR / "abstract.tex")

    bib = parse_bib(PAPER / "references.bib")

    # ---- Pass 1: numbering + citation order --------------------------------
    sec_map = {}
    fig_map = {}
    tab_map = {}
    cite_order = []

    def register_cites(text):
        for m in re.finditer(r"\\cite\{([^}]*)\}", text):
            for k in [x.strip() for x in m.group(1).split(",") if x.strip()]:
                if k not in cite_order:
                    cite_order.append(k)

    sec = sub = subsub = 0
    cur_number = None
    for name, blocks in all_blocks:
        for b in blocks:
            kind = b[0]
            if kind == "section":
                sec += 1
                sub = subsub = 0
                cur_number = str(sec)
            elif kind == "subsection":
                sub += 1
                subsub = 0
                cur_number = f"{sec}.{sub}"
            elif kind == "subsubsection":
                subsub += 1
                cur_number = f"{sec}.{sub}.{subsub}"
            elif kind == "label":
                if b[1].startswith("sec:"):
                    sec_map[b[1]] = cur_number
            elif kind == "figure":
                if b[3]:
                    fig_map[b[3]] = str(len(fig_map) + 1)
            elif kind == "table":
                _c, lab, _n, _h, _d = parse_table_tex(TABLES_DIR / (b[1].split("/")[-1] + ".tex"))
                if lab:
                    tab_map[lab] = str(len(tab_map) + 1)
            elif kind == "para":
                register_cites(b[1])
            elif kind == "enumerate":
                for it in b[1]:
                    register_cites(it)
    for b in abstract_blocks:
        if b[0] in ("abstract", "keywords"):
            register_cites(b[1])

    cite_map = {k: i + 1 for i, k in enumerate(cite_order)}
    ref_map = {**sec_map, **fig_map, **tab_map}

    # ---- Build document -----------------------------------------------------
    doc = Document()
    doc.core_properties.title = "Beyond In-Domain Accuracy: Cross-Dataset Generalization and Deployment Analysis of Object Detectors for Autonomous Driving"
    doc.core_properties.author = "Author Name"
    doc.core_properties.subject = "Cross-Dataset Object Detection"

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(1.0))

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    main_txt = re.sub(r"\s+", " ", (PAPER / "main.tex").read_text(encoding="utf-8"))
    title_text = extract_command(main_txt, "\\title")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title_text)
    tr.bold = True
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(16)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(12)
    ar = author_p.add_run("Author Name")
    ar.font.name = "Times New Roman"
    ar.font.size = Pt(12)

    def add_heading(text, level):
        p = doc.add_heading("", level=level)
        p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt({1: 14, 2: 12, 3: 11}[level])
        r.italic = level == 3
        return p

    def add_body(text, justify=True):
        p = doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for (t, b, it, code) in parse_inline(text, ref_map, cite_map):
            r = p.add_run(t)
            r.bold = b
            r.italic = it
            if code:
                r.font.name = "Consolas"
                r.font.size = Pt(10)
        return p

    def add_figure(figfile, caption, label):
        basename = figfile.split("/")[-1]
        png = png_map[basename]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(png), width=Inches(6.0))
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(8)
        run = cap_p.add_run("Figure " + label + ". ")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        for (t, b, it, code) in parse_inline(caption, ref_map, cite_map):
            r = cap_p.add_run(t)
            r.bold = b
            r.italic = it
            r.font.name = "Consolas" if code else "Times New Roman"
            r.font.size = Pt(10)

    # ---- Abstract + keywords ----------------------------------------------
    add_heading("Abstract", 1)
    for b in abstract_blocks:
        if b[0] == "abstract":
            add_body(b[1])
        elif b[0] == "keywords":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run("Index Terms\u2014 ")
            r.bold = True
            r.font.name = "Times New Roman"
            for (t, bd, itl, code) in parse_inline(b[1], ref_map, cite_map):
                r = p.add_run(t)
                r.bold = bd
                r.italic = itl

    # ---- Render numbered sections -----------------------------------------
    sec = sub = subsub = 0
    for name, blocks in all_blocks:
        for b in blocks:
            kind = b[0]
            if kind == "section":
                sec += 1
                sub = subsub = 0
                add_heading(f"{sec}. {b[1]}", 1)
            elif kind == "subsection":
                sub += 1
                subsub = 0
                add_heading(f"{sec}.{sub}. {b[1]}", 2)
            elif kind == "subsubsection":
                subsub += 1
                add_heading(f"{sec}.{sub}.{subsub}. {b[1]}", 3)
            elif kind == "para":
                add_body(b[1])
            elif kind == "enumerate":
                for idx, it in enumerate(b[1], 1):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(f"{idx}. ")
                    for (t, bd, itl, code) in parse_inline(it, ref_map, cite_map):
                        r = p.add_run(t)
                        r.bold = bd
                        r.italic = itl
                        if code:
                            r.font.name = "Consolas"
                            r.font.size = Pt(10)
            elif kind == "figure":
                _, img, cap, lab = b
                add_figure(img, cap, fig_map[lab])
            elif kind == "table":
                cap, lab, ncols, hrows, drows = parse_table_tex(TABLES_DIR / (b[1].split("/")[-1] + ".tex"))
                add_table(doc, cap, tab_map[lab], ncols, hrows, drows, ref_map, cite_map)

    # ---- References --------------------------------------------------------
    add_heading("References", 1)
    for i, key in enumerate(cite_order, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{i}] ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r2 = p.add_run(format_reference(bib[key]))
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)

    doc.save(OUTPUT)
    print("Saved", OUTPUT)


if __name__ == "__main__":
    main()
