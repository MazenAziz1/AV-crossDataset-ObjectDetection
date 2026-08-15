import json
from pathlib import Path

import pandas as pd
from docx import Document

DETECTOR_ORDER = ["yolo", "rtdetr", "retinanet", "faster_rcnn"]
NAMES = {
    "yolo": "YOLOv8s",
    "rtdetr": "RT-DETR-L",
    "retinanet": "RetinaNet",
    "faster_rcnn": "Faster R-CNN",
}
CLASSES = [("1", "Vehicle"), ("2", "Pedestrian"), ("3", "Cyclist")]


def r4(x):
    return f"{x:.4f}"


def r2(x):
    return f"{x:.2f}"


def main():
    root = Path(__file__).resolve().parents[2]
    kitti_dir = root / "outputs" / "milestone_4" / "metrics" / "kitti_validation"
    waymo_dir = root / "outputs" / "milestone_6" / "waymo_external_validation" / "metrics"
    ga_dir = root / "outputs" / "milestone_6" / "generalization_analysis"
    handoff_dir = root / "outputs" / "milestone_6" / "handoff_validation"

    kitti = {d: json.load(open(kitti_dir / f"{d}_metrics.json", encoding="utf-8")) for d in DETECTOR_ORDER}
    waymo = {d: json.load(open(waymo_dir / f"{d}_waymo_metrics.json", encoding="utf-8")) for d in DETECTOR_ORDER}

    comparison = pd.read_csv(ga_dir / "tables" / "kitti_vs_waymo_comparison.csv")
    ratio = pd.read_csv(ga_dir / "tables" / "generalization_ratio_table.csv")
    class_df = pd.read_csv(ga_dir / "tables" / "class_wise_degradation.csv")
    gen_summary = json.load(open(ga_dir / "generalization_analysis_summary.json", encoding="utf-8"))

    handoff = json.load(open(handoff_dir / "waymo_handoff_summary.json", encoding="utf-8"))

    comp_by_det = comparison.set_index("detector")

    # ---------- Document ----------
    doc = Document()
    doc.add_heading("Waymo External Validation and Generalization Analysis", level=0)
    sub = doc.add_paragraph("Milestone 6 Report \u2014 Cross-Dataset Object Detection")
    sub.runs[0].italic = True

    doc.add_heading("1. Milestone 6 Objective", level=1)
    doc.add_paragraph(
        "Milestone 6 evaluates how well the KITTI-trained detectors generalize to the Waymo Open Dataset "
        "without retraining. The four locked checkpoints (YOLOv8s, RT-DETR-L, RetinaNet, and Faster R-CNN) are "
        "applied directly to the frozen Waymo representative subset, and domain-shift degradation is measured by "
        "comparing the Milestone 4/5 KITTI validation metrics against the Waymo external validation metrics."
    )

    doc.add_heading("2. External Validation Methodology", level=1)
    doc.add_paragraph(
        "The Waymo external validation reuses the frozen Milestone 4/5 evaluation stack unchanged: predictions "
        "are produced by the Milestone 4 adapters (Ultralytics YOLO/RT-DETR and Torchvision RetinaNet/Faster "
        "R-CNN with the training-time ImageNet normalization), and metrics are computed with the same "
        "pycocotools COCOeval evaluator. The only change is the dataset: the ground truth and images are swapped "
        "from the KITTI validation partition (1,496 images) to the Waymo external subset (996 images). Per-image "
        "inference time is measured during evaluation to produce the mean inference time metric."
    )
    for line in [
        "Input size: 640 x 640.",
        "Confidence threshold for AP curves: 0.001 (fixed before evaluation).",
        "Operating point: confidence >= 0.25, IoU >= 0.50 (fixed before evaluation).",
        "No thresholds are tuned using Waymo results.",
        "Waymo has no DontCare-style ignore regions (Sign is excluded as a non-target class).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. Dataset and Waymo Subset Description", level=1)
    doc.add_paragraph(
        f"The Waymo external subset contains {handoff['checks']['coco_num_images']} images "
        f"({handoff['checks']['coco_num_annotations']} target boxes) from the Waymo Open Dataset validation "
        "split, camera FRONT, sampled every fifth front frame across 25 segments. The subset is frozen before "
        "model evaluation and is used only for external evaluation."
    )

    doc.add_heading("4. No-Retraining Policy", level=1)
    for line in [
        "No training, fine-tuning, hyperparameter selection, or checkpoint selection uses Waymo data.",
        "Thresholds are defined a priori in the frozen evaluation policy; none are tuned on Waymo results.",
        "All four models are the locked KITTI-trained checkpoints from Milestone 4/5.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("5. Model Checkpoints Used", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Detector"
    t.rows[0].cells[1].text = "Checkpoint"
    for d in DETECTOR_ORDER:
        rel = Path(waymo[d]["checkpoint"])
        try:
            rel = rel.relative_to(root)
        except ValueError:
            pass
        row = t.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = str(rel)

    doc.add_heading("6. Class Mapping", level=1)
    doc.add_paragraph(
        "The harmonized three-class task uses Vehicle, Pedestrian, and Cyclist. Canonical evaluation ids are the "
        "COCO category ids 1 (Vehicle), 2 (Pedestrian), and 3 (Cyclist). YOLO/RT-DETR internal class indices are "
        "0/1/2 and are mapped to COCO ids via +1; Torchvision labels 1/2/3 equal the COCO ids with label 0 as "
        "background. Waymo Sign boxes are excluded from the three-class task."
    )

    doc.add_heading("7. Evaluation Metrics", level=1)
    for line in [
        "mAP@0.50 (AP50)",
        "mAP@[0.50:0.95] (primary metric)",
        "Per-class AP50 and AP50-95",
        "Mean inference time (ms/image) on the real Waymo images",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("8. KITTI Baseline Summary", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    for i, s in enumerate(["Detector", "mAP@0.50", "mAP@0.50:0.95", "Vehicle AP", "Pedestrian AP", "Cyclist AP"]):
        h[i].text = s
    for d in DETECTOR_ORDER:
        m = kitti[d]
        pc = m["per_class"]
        row = t.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = r4(m["metrics"]["mAP_50"])
        row[2].text = r4(m["metrics"]["mAP_50_95"])
        row[3].text = r4(pc["1"]["AP_50_95"])
        row[4].text = r4(pc["2"]["AP_50_95"])
        row[5].text = r4(pc["3"]["AP_50_95"])

    doc.add_heading("9. Waymo External Validation Results", level=1)
    t = doc.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    for i, s in enumerate(["Detector", "mAP@0.50", "mAP@0.50:0.95", "Vehicle AP", "Pedestrian AP", "Cyclist AP", "Mean inf. (ms)"]):
        h[i].text = s
    for d in DETECTOR_ORDER:
        m = waymo[d]
        pc = m["per_class"]
        row = t.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = r4(m["metrics"]["mAP_50"])
        row[2].text = r4(m["metrics"]["mAP_50_95"])
        row[3].text = r4(pc["1"]["AP_50_95"])
        row[4].text = r4(pc["2"]["AP_50_95"])
        row[5].text = r4(pc["3"]["AP_50_95"])
        row[6].text = r2(m["mean_inference_ms"])

    doc.add_heading("10. KITTI vs Waymo Comparison", level=1)
    t = doc.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    for i, s in enumerate(["Detector", "KITTI mAP@0.50:0.95", "Waymo mAP@0.50:0.95", "Absolute drop", "Drop %", "Generalization ratio", "Rank"]):
        h[i].text = s
    for d in DETECTOR_ORDER:
        row = t.add_row().cells
        row[0].text = NAMES[d]
        row[1].text = r4(comp_by_det.loc[d, "KITTI_mAP50_95"])
        row[2].text = r4(comp_by_det.loc[d, "Waymo_mAP50_95"])
        row[3].text = r4(comp_by_det.loc[d, "mAP50_95_absolute_drop"])
        row[4].text = r2(comp_by_det.loc[d, "mAP50_95_drop_percent"])
        row[5].text = r4(comp_by_det.loc[d, "mAP50_95_generalization_ratio"])
        row[6].text = str(int(comp_by_det.loc[d, "rank_by_mAP50_95_generalization"]))

    doc.add_heading("11. Generalization Ratio Analysis", level=1)
    doc.add_paragraph(
        f"The Generalization Ratio (Waymo mAP / KITTI mAP) quantifies cross-domain stability. "
        f"The best generalizing detector is {NAMES[gen_summary['best_generalizing_detector']]} "
        f"(mAP@0.50:0.95 ratio {r4(comp_by_det.loc[gen_summary['best_generalizing_detector'], 'mAP50_95_generalization_ratio'])}), "
        f"and the worst is {NAMES[gen_summary['worst_generalizing_detector']]} "
        f"(ratio {r4(comp_by_det.loc[gen_summary['worst_generalizing_detector'], 'mAP50_95_generalization_ratio'])})."
    )

    doc.add_heading("12. Class-Wise Degradation Analysis", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    for i, s in enumerate(["Detector", "Class", "KITTI AP@0.50:0.95", "Waymo AP@0.50:0.95", "Absolute drop", "Drop %"]):
        h[i].text = s
    for d in DETECTOR_ORDER:
        sub_df = class_df[class_df["detector"] == d]
        for _, r in sub_df.iterrows():
            row = t.add_row().cells
            row[0].text = NAMES[d]
            row[1].text = r["class_name"]
            row[2].text = r4(r["KITTI_AP50_95"])
            row[3].text = r4(r["Waymo_AP50_95"])
            row[4].text = r4(r["absolute_drop"])
            row[5].text = r2(r["drop_percent"])
    doc.add_paragraph(
        f"The class with the largest mean absolute drop across detectors is "
        f"{gen_summary['largest_degraded_class']}."
    )

    doc.add_heading("13. Domain-Shift Discussion", level=1)
    doc.add_paragraph(
        "All four detectors experience a severe accuracy drop when transferred from KITTI to Waymo, with "
        "mAP@0.50:0.95 falling by roughly 80-90% relative to the in-domain baseline. This is expected for a "
        "direct, no-retraining cross-dataset transfer: KITTI and Waymo differ in camera calibration, sensor "
        "position, scene composition, object scale distribution, and annotation density. The ranking of "
        "detectors inverts across domains: YOLOv8s, the strongest in-domain detector, shows the lowest "
        "generalization ratio, while RetinaNet, weaker in-domain, generalizes best. This inversion is a "
        "candidate indicator of overfitting to KITTI-specific statistics for the higher-capacity detectors."
    )

    doc.add_heading("14. Paper-Ready External Validation Subsection", level=1)
    doc.add_paragraph(
        "Cross-dataset generalization was assessed by applying the four KITTI-trained detectors directly to a "
        "frozen Waymo representative subset (996 front-camera images) with no retraining, fine-tuning, or "
        "threshold tuning. All detectors degrade sharply on Waymo: mAP@[0.50:0.95] drops from 0.538-0.690 "
        "(KITTI) to 0.065-0.097 (Waymo), corresponding to generalization ratios of 0.095-0.180. RetinaNet is the "
        "most domain-stable detector (ratio 0.180) and YOLOv8s the least (ratio 0.095). Vehicle is the class with "
        "the largest degradation across detectors. These results indicate that the in-domain ranking does not "
        "transfer to the out-of-domain setting, motivating domain-robust training or adaptation strategies."
    )

    doc.add_heading("15. Limitations and Threats to Validity", level=1)
    for line in [
        "The Waymo subset (996 images) is smaller than the KITTI validation split (1,496 images).",
        "Inference time is measured per real Waymo image including the first-image warmup; it is not directly "
        "comparable to the dummy-input benchmark latency reported in Milestone 4.",
        "No confidence threshold is tuned on Waymo, so operating-point precision/recall reflect the frozen 0.25 "
        "threshold, not an optimized one.",
        "Results are specific to the locked checkpoints and the frozen subset; they should not be generalized "
        "beyond this configuration.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("16. Generated Artifacts", level=1)
    for line in [
        "Handoff validation: outputs/milestone_6/handoff_validation/",
        "Waymo metrics: outputs/milestone_6/waymo_external_validation/metrics/",
        "Comparison tables: outputs/milestone_6/generalization_analysis/tables/",
        "Figures: outputs/milestone_6/figures/",
        "Final audit: outputs/milestone_6/final_audit/",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    out_docx = root / "docs" / "milestone_6" / "Milestone_6_Waymo_External_Validation_Report.docx"
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print(f"Saved: {out_docx}")

    # ---------- Markdown ----------
    md = []
    md.append("# Milestone 6 \u2014 Waymo External Validation and Generalization Analysis")
    md.append("")
    md.append("## 1. Milestone 6 Objective")
    md.append("")
    md.append("Evaluate how well the KITTI-trained detectors generalize to Waymo without retraining, by applying "
              "the four locked checkpoints directly to the frozen Waymo representative subset and measuring "
              "domain-shift degradation against the Milestone 4/5 KITTI validation baseline.")
    md.append("")
    md.append("## 2. External Validation Methodology")
    md.append("")
    md.append("The Waymo evaluation reuses the frozen Milestone 4/5 evaluation stack (adapters + pycocotools "
              "COCOeval). Input size 640x640, AP curve confidence 0.001, operating point confidence 0.25 / IoU 0.50, "
              "no threshold tuning on Waymo, and no DontCare-style ignore regions (Sign excluded).")
    md.append("")
    md.append("## 3. Dataset and Waymo Subset Description")
    md.append("")
    md.append(f"- Images: {handoff['checks']['coco_num_images']}")
    md.append(f"- Target boxes: {handoff['checks']['coco_num_annotations']}")
    md.append("- Source: Waymo Open Dataset validation split, FRONT camera, every 5th frame, 25 segments.")
    md.append("")
    md.append("## 4. No-Retraining Policy")
    md.append("")
    md.append("No training/fine-tuning/hyperparameter/checkpoint/threshold selection uses Waymo data.")
    md.append("")
    md.append("## 5. Model Checkpoints Used")
    md.append("")
    for d in DETECTOR_ORDER:
        md.append(f"- {NAMES[d]}: {waymo[d]['checkpoint']}")
    md.append("")
    md.append("## 6. Class Mapping")
    md.append("")
    md.append("| COCO id | YOLO id | Torchvision label | Name |")
    md.append("|---|---|---|---|")
    md.append("| 1 | 0 | 1 | Vehicle |")
    md.append("| 2 | 1 | 2 | Pedestrian |")
    md.append("| 3 | 2 | 3 | Cyclist |")
    md.append("")
    md.append("## 7. Evaluation Metrics")
    md.append("")
    md.append("mAP@0.50, mAP@[0.50:0.95], per-class AP50/AP50-95, mean inference time.")
    md.append("")
    md.append("## 8. KITTI Baseline Summary")
    md.append("")
    md.append("| Detector | mAP@0.50 | mAP@0.50:0.95 | Vehicle AP | Pedestrian AP | Cyclist AP |")
    md.append("|---|---|---|---|---|---|")
    for d in DETECTOR_ORDER:
        m = kitti[d]; pc = m["per_class"]
        md.append(f"| {NAMES[d]} | {r4(m['metrics']['mAP_50'])} | {r4(m['metrics']['mAP_50_95'])} | "
                  f"{r4(pc['1']['AP_50_95'])} | {r4(pc['2']['AP_50_95'])} | {r4(pc['3']['AP_50_95'])} |")
    md.append("")
    md.append("## 9. Waymo External Validation Results")
    md.append("")
    md.append("| Detector | mAP@0.50 | mAP@0.50:0.95 | Vehicle AP | Pedestrian AP | Cyclist AP | Mean inf. (ms) |")
    md.append("|---|---|---|---|---|---|---|")
    for d in DETECTOR_ORDER:
        m = waymo[d]; pc = m["per_class"]
        md.append(f"| {NAMES[d]} | {r4(m['metrics']['mAP_50'])} | {r4(m['metrics']['mAP_50_95'])} | "
                  f"{r4(pc['1']['AP_50_95'])} | {r4(pc['2']['AP_50_95'])} | {r4(pc['3']['AP_50_95'])} | "
                  f"{r2(m['mean_inference_ms'])} |")
    md.append("")
    md.append("## 10. KITTI vs Waymo Comparison")
    md.append("")
    md.append("| Detector | KITTI mAP@0.50:0.95 | Waymo mAP@0.50:0.95 | Absolute drop | Drop % | Ratio | Rank |")
    md.append("|---|---|---|---|---|---|---|")
    for d in DETECTOR_ORDER:
        c = comp_by_det.loc[d]
        md.append(f"| {NAMES[d]} | {r4(c['KITTI_mAP50_95'])} | {r4(c['Waymo_mAP50_95'])} | "
                  f"{r4(c['mAP50_95_absolute_drop'])} | {r2(c['mAP50_95_drop_percent'])} | "
                  f"{r4(c['mAP50_95_generalization_ratio'])} | {int(c['rank_by_mAP50_95_generalization'])} |")
    md.append("")
    md.append("## 11. Generalization Ratio Analysis")
    md.append("")
    md.append(f"- Best generalizing detector: **{NAMES[gen_summary['best_generalizing_detector']]}**")
    md.append(f"- Worst generalizing detector: **{NAMES[gen_summary['worst_generalizing_detector']]}**")
    md.append("")
    md.append("## 12. Class-Wise Degradation Analysis")
    md.append("")
    md.append(f"- Largest degraded class: **{gen_summary['largest_degraded_class']}**")
    md.append("")
    md.append("## 13. Domain-Shift Discussion")
    md.append("")
    md.append("All detectors degrade sharply (80-90% relative mAP@0.50:0.95 drop). The in-domain detector ranking "
              "inverts on Waymo: YOLOv8s is strongest on KITTI but least generalizing, while RetinaNet is weaker "
              "in-domain yet most stable out-of-domain, suggesting domain overfitting in the higher-capacity "
              "detectors.")
    md.append("")
    md.append("## 14. Paper-Ready External Validation Subsection")
    md.append("")
    md.append("Cross-dataset generalization was assessed by applying the four KITTI-trained detectors directly to a "
              "frozen Waymo representative subset (996 front-camera images) with no retraining, fine-tuning, or "
              "threshold tuning. All detectors degrade sharply on Waymo: mAP@[0.50:0.95] drops from 0.538-0.690 "
              "(KITTI) to 0.065-0.097 (Waymo), corresponding to generalization ratios of 0.095-0.180. RetinaNet is "
              "the most domain-stable detector (ratio 0.180) and YOLOv8s the least (ratio 0.095). Vehicle is the "
              "class with the largest degradation across detectors. These results indicate that the in-domain "
              "ranking does not transfer to the out-of-domain setting, motivating domain-robust training or "
              "adaptation strategies.")
    md.append("")
    md.append("## 15. Limitations and Threats to Validity")
    md.append("")
    md.append("- Waymo subset (996) is smaller than KITTI validation (1,496).")
    md.append("- Inference time includes first-image warmup and is not comparable to Milestone 4 dummy-input latency.")
    md.append("- Operating point uses the frozen 0.25 confidence threshold (not optimized on Waymo).")
    md.append("- Results are specific to the locked checkpoints and frozen subset.")
    md.append("")
    md.append("## 16. Generated Artifacts")
    md.append("")
    md.append("- Handoff validation: outputs/milestone_6/handoff_validation/")
    md.append("- Waymo metrics: outputs/milestone_6/waymo_external_validation/metrics/")
    md.append("- Comparison tables: outputs/milestone_6/generalization_analysis/tables/")
    md.append("- Figures: outputs/milestone_6/figures/")
    md.append("- Final audit: outputs/milestone_6/final_audit/")
    md.append("")

    out_md = root / "docs" / "milestone_6" / "Milestone_6_Waymo_External_Validation_Report.md"
    out_md.write_text("\n".join(md), encoding="utf-8")
    print(f"Saved: {out_md}")


if __name__ == "__main__":
    main()
